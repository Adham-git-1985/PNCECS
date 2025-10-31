# -*- coding: utf-8 -*-
"""
build_poetry_docx_align.py (page size + A5)
------------------------------------------
- الشطر الأول على اليمين، الثاني على اليسار (مع خيارات محاذاة لكل شطر).
- تخطيط: صف واحد (--layout row) أو صفّان لكل بيت (--layout two-rows).
- فراغ بين الشطرين (--gap) بأسلوبين: --gap-mode column | indent.
- أبعاد تلقائية للأعمدة (--auto-widths) محسوبة من الصفحة وهوامشها.
- جديد: تغيير مقاس الصفحة وهوامشها:
    --page {A4, A5, Letter}  (A5 = 14.8 × 21.0 سم)
    --margin 2.0             (سم لكل الحواف؛ يمكن تغيير القيمة)

مثال يطابق طلبك (A5 + كل شطر بسطر منفصل + الشطر الأول يمين/الثاني يسار):
    python build_poetry_docx_align.py --in poems.txt --out ديواني.docx --page A5 --margin 2 \
      --layout two-rows --auto-widths --gap 1.0 --gap-mode column \
      --right-align right --second-align left --font "Traditional Arabic" --size 20
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys


# ----------------- Helpers -----------------
def set_paragraph_rtl(paragraph, enable=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for tag, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn('w:w'), str(int(val)))
        el.set(qn('w:type'), 'dxa')


def set_table_borders(table, show=False):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn('w:val'), 'single' if show else 'nil')


def set_table_direction_rtl(table, enable=True):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bidi = tblPr.find(qn('w:bidiVisual'))
    if bidi is None:
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')


def read_poems_from_file(path):
    poems = []
    for raw in Path(path).read_text(encoding='utf-8').splitlines():
        line = raw.strip()
        if not line:
            continue
        if '|' in line:
            a, b = line.split('|', 1)
        elif '\t' in line:
            a, b = line.split('\t', 1)
        else:
            a, b = line, ''
        poems.append((a.strip(), b.strip()))
    return poems


# ----------------- Alignment mapping -----------------
def map_align(value, default='right'):
    val = (value or default).lower()
    if val == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    if val == 'left':
        return WD_ALIGN_PARAGRAPH.LEFT
    if val == 'justify':
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if val == 'distribute':
        return WD_ALIGN_PARAGRAPH.DISTRIBUTE
    return WD_ALIGN_PARAGRAPH.RIGHT


# ----------------- Page size/margins -----------------
def apply_page_setup(doc, page, margin_cm):
    """Set page size and uniform margins (cm)."""
    sec = doc.sections[0]
    page = (page or 'A4').upper()
    if page == 'A5':
        sec.page_width  = Cm(14.8)
        sec.page_height = Cm(21.0)
    elif page == 'LETTER':
        # 8.5 × 11 inches
        sec.page_width  = Cm(21.59)
        sec.page_height = Cm(27.94)
    else:  # A4 default
        sec.page_width  = Cm(21.0)
        sec.page_height = Cm(29.7)
    if margin_cm is not None:
        m = Cm(margin_cm)
        sec.left_margin = m
        sec.right_margin = m
        sec.top_margin = m
        sec.bottom_margin = m


def compute_auto_widths(doc, gap_cm, gap_mode):
    sec = doc.sections[0]
    available = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    available = max(0.1, available)
    if gap_mode == 'column':
        inner = max(0.1, available - max(0.0, gap_cm))
        each = max(0.1, inner / 2.0)
        return each, each
    else:
        each = max(0.1, available / 2.0)
        return each, each


# ----------------- Core -----------------
def build_poetry_docx_align(
    poems, out_path="poetry_aligned.docx",
    col_width_cm=(7.5, 7.5), gap_cm=0.0, gap_mode="column",
    font_name="Traditional Arabic", font_size_pt=18,
    show_borders=False, rtl=True,
    right_align_mode='right', second_align_mode='left',
    auto_widths=False, layout='row',
    page='A4', margin=2.0
):
    """layout: 'row' أو 'two-rows'. page: A4/A5/Letter. margin: سم لكل الحواف."""
    if gap_mode not in ("indent", "column"):
        raise ValueError("gap_mode must be 'indent' or 'column'")
    if layout not in ("row", "two-rows"):
        raise ValueError("layout must be 'row' or 'two-rows'")

    doc = Document()

    # 1) Apply page size and margins
    apply_page_setup(doc, page, margin)

    # 2) Auto widths AFTER page/margins have been set
    if auto_widths:
        w1, w2 = compute_auto_widths(doc, gap_cm, gap_mode)
    else:
        w1, w2 = col_width_cm

    align_first  = map_align(right_align_mode, default='right')
    align_second = map_align(second_align_mode, default='left')

    # Build table
    if layout == "row":
        if gap_mode == "column":
            table = doc.add_table(rows=len(poems), cols=3)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(max(0.0, gap_cm))
            table.columns[2].width = Cm(w2)
            set_table_borders(table, show=show_borders)
            set_table_direction_rtl(table, True)
            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                row = table.rows[i]
                c_r = row.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                p_r = c_r.paragraphs[0]
                p_r.alignment = align_first
                set_paragraph_rtl(p_r, rtl)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

                row.cells[1].width = Cm(max(0.0, gap_cm))

                c_l = row.cells[2]; c_l.width = Cm(w2)
                set_cell_margins(c_l, 0, 0, 0, 0)
                p_l = c_l.paragraphs[0]
                p_l.alignment = align_second
                set_paragraph_rtl(p_l, rtl)
                run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)
        else:
            table = doc.add_table(rows=len(poems), cols=2)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(w2)
            set_table_borders(table, show=show_borders)
            set_table_direction_rtl(table, True)
            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                row = table.rows[i]
                c_r = row.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                p_r = c_r.paragraphs[0]
                p_r.alignment = align_first
                set_paragraph_rtl(p_r, rtl)
                if gap_cm > 0:
                    p_r.paragraph_format.right_indent = Cm(gap_cm)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

                c_l = row.cells[1]; c_l.width = Cm(w2)
                set_cell_margins(c_l, 0, 0, 0, 0)
                p_l = c_l.paragraphs[0]
                p_l.alignment = align_second
                set_paragraph_rtl(p_l, rtl)
                if gap_cm > 0:
                    p_l.paragraph_format.left_indent = Cm(gap_cm)
                run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

    else:  # two-rows
        if gap_mode == "column":
            table = doc.add_table(rows=len(poems)*2, cols=3)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(max(0.0, gap_cm))
            table.columns[2].width = Cm(w2)
            set_table_borders(table, show=show_borders)
            set_table_direction_rtl(table, True)
            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                r1 = table.rows[i*2]
                r2 = table.rows[i*2 + 1]

                c_r = r1.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                p_r = c_r.paragraphs[0]
                p_r.alignment = align_first
                set_paragraph_rtl(p_r, rtl)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

                r1.cells[1].width = Cm(max(0.0, gap_cm))

                c_l = r2.cells[2]; c_l.width = Cm(w2)
                set_cell_margins(c_l, 0, 0, 0, 0)
                p_l = c_l.paragraphs[0]
                p_l.alignment = align_second
                set_paragraph_rtl(p_l, rtl)
                run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)
        else:
            table = doc.add_table(rows=len(poems)*2, cols=2)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(w2)
            set_table_borders(table, show=show_borders)
            set_table_direction_rtl(table, True)
            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                r1 = table.rows[i*2]
                r2 = table.rows[i*2 + 1]

                c_r = r1.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                p_r = c_r.paragraphs[0]
                p_r.alignment = align_first
                set_paragraph_rtl(p_r, rtl)
                if gap_cm > 0:
                    p_r.paragraph_format.right_indent = Cm(gap_cm)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

                c_l = r2.cells[1]; c_l.width = Cm(w2)
                set_cell_margins(c_l, 0, 0, 0, 0)
                p_l = c_l.paragraphs[0]
                p_l.alignment = align_second
                set_paragraph_rtl(p_l, rtl)
                if gap_cm > 0:
                    p_l.paragraph_format.left_indent = Cm(gap_cm)
                run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Arabic poetry layout with page size/margins and A5 support.")
    parser.add_argument("--in", dest="infile", default=None, help="UTF-8 file: each line is a verse: first|second or Tab")
    parser.add_argument("--out", dest="outfile", default="poetry_aligned.docx", help="Output .docx filename")
    parser.add_argument("--page", choices=["A4","A5","Letter"], default="A4", help="Page size")
    parser.add_argument("--margin", type=float, default=2.0, help="Uniform page margins (cm)")
    parser.add_argument("--w1", type=float, default=7.5, help="Right column width (cm)")
    parser.add_argument("--w2", type=float, default=7.5, help="Left column width (cm)")
    parser.add_argument("--gap", type=float, default=0.0, help="Gap between hemistichs (cm)")
    parser.add_argument("--gap-mode", choices=["indent","column"], default="column", help="Gap as indent or third column")
    parser.add_argument("--font", type=str, default="Traditional Arabic", help="Font family")
    parser.add_argument("--size", type=float, default=18, help="Font size (pt)")
    parser.add_argument("--borders", action="store_true", help="Show table borders")
    parser.add_argument("--rtl-off", action="store_true", help="Disable RTL (default on)")
    parser.add_argument("--right-align", choices=["right","justify","distribute"], default="right",
                        help="Alignment for first hemistich")
    parser.add_argument("--second-align", choices=["left","justify","distribute"], default="left",
                        help="Alignment for second hemistich")
    parser.add_argument("--auto-widths", action="store_true", help="Compute w1/w2 from page width & margins automatically.")
    parser.add_argument("--layout", choices=["row","two-rows"], default="row", help="Place both hemistichs in the same row or each in its own row.")
    args = parser.parse_args()

    # Load poems
    if args.infile:
        poems = read_poems_from_file(args.infile)
        if not poems:
            print("⚠️ الملف لا يحتوي أبياتًا صالحة.", file=sys.stderr); sys.exit(1)
    else:
        poems = [
            ("بوركت يا بلدي فيك الوفاء مَثَلُ", "يا شامَةً تُوِّجَتْ وطْني فلسطيننا"),
            ("ها جئتَ يا إخوتي أهدي فؤادي لي", "أشعُّ بالنور من جبلٍ إلى جبل"),
        ]

    out = build_poetry_docx_align(
        poems,
        out_path=args.outfile,
        col_width_cm=(args.w1, args.w2),
        gap_cm=args.gap,
        gap_mode=args.gap_mode,
        font_name=args.font,
        font_size_pt=args.size,
        show_borders=args.borders,
        rtl=(not args.rtl_off),
        right_align_mode=args.right_align,
        second_align_mode=args.second_align,
        auto_widths=args.auto_widths,
        layout=args.layout,
        page=args.page,
        margin=args.margin
    )
    print("✅ Saved ->", out)


if __name__ == "__main__":
    main()
