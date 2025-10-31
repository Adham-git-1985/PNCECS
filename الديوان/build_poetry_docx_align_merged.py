# -*- coding: utf-8 -*-
"""
build_poetry_docx_align_merged.py
---------------------------------
يحافظ على سلوك تنسيق الدواوين كما هو، ويضيف ميزات اختيارية:
  1) --merge-second        : دمج خلايا الصف الثاني (شطر البيت الثاني) في خلية واحدة.
  2) --second-shift-pct N  : ترك فراغ أبيض يسار الشطر الثاني بنسبة N%% من عرض عمود الشطر الأول (w1).
  3) --equalize-lengths    : اجعل مساحة نص الشطر الثاني مساوية تمامًا لمساحة نص الشطر الأول (w1).
     - تعمل الإضافات مع --layout two-rows تحديدًا، وتحترم --second-align (right/left/justify/distribute).

مثال تشغيل:
    python build_poetry_docx_align_merged.py --in poems.txt --out ديواني.docx --page A5 --margin 1 \
      --layout two-rows --auto-widths --gap 0.05 --gap-mode indent \
      --right-align distribute --second-align distribute \
      --font "Sakkal Majalla" --size 18 --nowrap \
      --merge-second --equalize-lengths
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys

# ----------------- Helpers -----------------
def set_paragraph_rtl(paragraph, enable=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')

def set_table_borders(table, show=False):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn('w:val'), 'single' if show else 'nil')

def set_table_direction_rtl(table, enable=True):
    tblPr = table._tbl.tblPr
    bidi = tblPr.find(qn('w:bidiVisual'))
    if bidi is None:
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')

def disable_cell_wrap(cell):
    """منع التفاف النص داخل الخلية (يبقى سطر واحد)."""
    tcPr = cell._tc.get_or_add_tcPr()
    no_wrap = tcPr.find(qn('w:noWrap'))
    if no_wrap is None:
        no_wrap = OxmlElement('w:noWrap')
        tcPr.append(no_wrap)

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """يضبط هوامش الخلية بوحدات twips (1 سم ≈ 566.93 twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for tag, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}"); tcMar.append(el)
        el.set(qn('w:w'), str(int(val))); el.set(qn('w:type'), 'dxa')

def cm_to_twips(cm):
    return int(round(cm * 566.9291339))

def read_poems_from_file(path):
    poems = []
    for raw in Path(path).read_text(encoding='utf-8').splitlines():
        line = raw.strip()
        if not line:
            continue
        if '|' in line:
            a, b = line.split('|', 1)
        elif '\t' in line:
            a, b = line.split('\t', 1)
        else:
            a, b = line, ''
        poems.append((a.strip(), b.strip()))
    return poems

# ----------------- Alignment mapping -----------------
def map_align(value, default='right'):
    val = (value or default).lower()
    if val == 'right': return WD_ALIGN_PARAGRAPH.RIGHT
    if val == 'left': return WD_ALIGN_PARAGRAPH.LEFT
    if val == 'justify': return WD_ALIGN_PARAGRAPH.JUSTIFY
    if val == 'distribute': return WD_ALIGN_PARAGRAPH.DISTRIBUTE
    return WD_ALIGN_PARAGRAPH.RIGHT

# ----------------- Page size/margins -----------------
def apply_page_setup(doc, page, margin_cm):
    sec = doc.sections[0]
    page = (page or 'A4').upper()
    if page == 'A5':
        sec.page_width  = Cm(14.8); sec.page_height = Cm(21.0)
    elif page == 'LETTER':
        sec.page_width  = Cm(21.59); sec.page_height = Cm(27.94)
    else:  # A4
        sec.page_width  = Cm(21.0);  sec.page_height = Cm(29.7)
    if margin_cm is not None:
        m = Cm(margin_cm)
        sec.left_margin = m; sec.right_margin = m; sec.top_margin = m; sec.bottom_margin = m

def compute_auto_widths(doc, gap_cm, gap_mode):
    sec = doc.sections[0]
    available = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    available = max(0.1, available)
    if gap_mode == 'column':
        inner = max(0.1, available - max(0.0, gap_cm))
        each = max(0.1, inner / 2.0)
        return each, each
    else:
        each = max(0.1, available / 2.0)
        return each, each

# ----------------- Core -----------------
def build_poetry_docx_align(
    poems, out_path="poetry_aligned.docx",
    col_width_cm=(7.5, 7.5), gap_cm=0.0, gap_mode="column",
    font_name="Traditional Arabic", font_size_pt=18,
    show_borders=False, rtl=True,
    right_align_mode='right', second_align_mode='left',
    auto_widths=False, layout='row',
    page='A4', margin=2.0, nowrap=False,
    merge_second=False, second_shift_pct=0.0, equalize_lengths=False,
):
    """
    يحافظ على السلوك القديم إذا لم تُفعّل الخيارات الجديدة.
    - merge_second + second_shift_pct / equalize_lengths تعمل فقط مع layout='two-rows'.
    """
    # تحقق قيم الخيارات
    if gap_mode not in ("indent", "column"):
        raise ValueError("gap_mode must be 'indent' or 'column'")
    if layout not in ("row", "two-rows"):
        raise ValueError("layout must be 'row' or 'two-rows'")

    doc = Document()
    apply_page_setup(doc, page, margin)

    # عرض الأعمدة
    if auto_widths:
        w1, w2 = compute_auto_widths(doc, gap_cm, gap_mode)
    else:
        w1, w2 = col_width_cm

    align_first  = map_align(right_align_mode, default='right')
    align_second = map_align(second_align_mode, default='left')

    # مقدار الإزاحة كنسبة من w1 (يُستخدم فقط عند الدمج)
    shift_cm = 0.0
    if layout == 'two-rows' and merge_second and second_shift_pct:
        shift_cm = max(0.0, (second_shift_pct or 0.0) * 0.01 * w1)

    if layout == "row":
        # === سلوك الصف الواحد كما هو ===
        if gap_mode == "column":
            table = doc.add_table(rows=len(poems), cols=3)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(max(0.0, gap_cm))
            table.columns[2].width = Cm(w2)
            set_table_borders(table, show=show_borders); set_table_direction_rtl(table, True)

            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                row = table.rows[i]
                # الشطر الأول
                c_r = row.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                if nowrap: disable_cell_wrap(c_r)
                p_r = c_r.paragraphs[0]; p_r.alignment = align_first; set_paragraph_rtl(p_r, rtl)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)
                # فجوة
                row.cells[1].width = Cm(max(0.0, gap_cm))
                # الشطر الثاني
                c_l = row.cells[2]; c_l.width = Cm(w2)
                set_cell_margins(c_l, 0, 0, 0, 0)
                if nowrap: disable_cell_wrap(c_l)
                p_l = c_l.paragraphs[0]; p_l.alignment = align_second; set_paragraph_rtl(p_l, rtl)
                run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)
        else:
            table = doc.add_table(rows=len(poems), cols=2)
            table.autofit = False
            table.columns[0].width = Cm(w1); table.columns[1].width = Cm(w2)
            set_table_borders(table, show=show_borders); set_table_direction_rtl(table, True)

            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                row = table.rows[i]
                # الشطر الأول
                c_r = row.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                if nowrap: disable_cell_wrap(c_r)
                p_r = c_r.paragraphs[0]; p_r.alignment = align_first; set_paragraph_rtl(p_r, rtl)
                if gap_cm > 0: p_r.paragraph_format.right_indent = Cm(gap_cm)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)
                # الشطر الثاني
                c_l = row.cells[1]; c_l.width = Cm(w2)
                set_cell_margins(c_l, 0, 0, 0, 0)
                if nowrap: disable_cell_wrap(c_l)
                p_l = c_l.paragraphs[0]; p_l.alignment = align_second; set_paragraph_rtl(p_l, rtl)
                if gap_cm > 0: p_l.paragraph_format.left_indent = Cm(gap_cm)
                run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

    else:
        # === وضع صفّين لكل بيت ===
        if gap_mode == "column":
            table = doc.add_table(rows=len(poems)*2, cols=3)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(max(0.0, gap_cm))
            table.columns[2].width = Cm(w2)
            set_table_borders(table, show=show_borders); set_table_direction_rtl(table, True)

            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                r1 = table.rows[i*2]; r2 = table.rows[i*2+1]
                # الشطر الأول (صف علوي، عمود 0)
                c_r = r1.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                if nowrap: disable_cell_wrap(c_r)
                p_r = c_r.paragraphs[0]; p_r.alignment = align_first; set_paragraph_rtl(p_r, rtl)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

                # الفاصل العمودي
                r1.cells[1].width = Cm(max(0.0, gap_cm))

                if merge_second:
                    # دمج خلايا الصف الثاني كلها: 0+1+2
                    merged = r2.cells[0].merge(r2.cells[1]).merge(r2.cells[2])
                    p2 = merged.paragraphs[0]; p2.alignment = align_second; set_paragraph_rtl(p2, rtl)

                    # مساواة الطول أم إزاحة بالنسبة المئوية؟
                    if equalize_lengths:
                        # إجمالي عرض الصف الثاني = w1 + gap + w2
                        # نريد مساحة النص = w1 → فراغ يسار = gap + w2
                        end_cm = max(0.0, max(0.0, gap_cm) + w2)
                    else:
                        end_cm = max(0.0, shift_cm)

                    set_cell_margins(merged, top=0, start=0, bottom=0, end=cm_to_twips(end_cm))
                    if nowrap: disable_cell_wrap(merged)
                    run2 = p2.add_run(left_hemistich); run2.font.name = font_name; run2.font.size = Pt(font_size_pt)
                else:
                    # السلوك السابق: الشطر الثاني في عمود 2 من الصف الثاني
                    c_l = r2.cells[2]; c_l.width = Cm(w2)
                    set_cell_margins(c_l, 0, 0, 0, 0)
                    if nowrap: disable_cell_wrap(c_l)
                    p_l = c_l.paragraphs[0]; p_l.alignment = align_second; set_paragraph_rtl(p_l, rtl)
                    run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

        else:
            table = doc.add_table(rows=len(poems)*2, cols=2)
            table.autofit = False
            table.columns[0].width = Cm(w1)
            table.columns[1].width = Cm(w2)
            set_table_borders(table, show=show_borders); set_table_direction_rtl(table, True)

            for i, (right_hemistich, left_hemistich) in enumerate(poems):
                r1 = table.rows[i*2]; r2 = table.rows[i*2+1]

                # الشطر الأول
                c_r = r1.cells[0]; c_r.width = Cm(w1)
                set_cell_margins(c_r, 0, 0, 0, 0)
                if nowrap: disable_cell_wrap(c_r)
                p_r = c_r.paragraphs[0]; p_r.alignment = align_first; set_paragraph_rtl(p_r, rtl)
                if gap_cm > 0: p_r.paragraph_format.right_indent = Cm(gap_cm)
                run_r = p_r.add_run(right_hemistich); run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

                if merge_second:
                    # دمج الخليتين في الصف الثاني (عمود0 + عمود1)
                    merged = r2.cells[0].merge(r2.cells[1])
                    p2 = merged.paragraphs[0]; p2.alignment = align_second; set_paragraph_rtl(p2, rtl)

                    # مساواة الطول أم إزاحة تراكمية (gap + shift)؟
                    if equalize_lengths:
                        # إجمالي العرض = w1 + w2 ؛ نريد مساحة النص = w1 → فراغ يسار = w2
                        end_cm = max(0.0, w2)
                    else:
                        end_cm = max(0.0, max(0.0, gap_cm) + max(0.0, shift_cm))

                    set_cell_margins(merged, top=0, start=0, bottom=0, end=cm_to_twips(end_cm))
                    if nowrap: disable_cell_wrap(merged)
                    run2 = p2.add_run(left_hemistich); run2.font.name = font_name; run2.font.size = Pt(font_size_pt)
                else:
                    # السلوك السابق: العمود الثاني مع left_indent = gap
                    c_l = r2.cells[1]; c_l.width = Cm(w2)
                    set_cell_margins(c_l, 0, 0, 0, 0)
                    if nowrap: disable_cell_wrap(c_l)
                    p_l = c_l.paragraphs[0]; p_l.alignment = align_second; set_paragraph_rtl(p_l, rtl)
                    if gap_cm > 0: p_l.paragraph_format.left_indent = Cm(gap_cm)
                    run_l = p_l.add_run(left_hemistich); run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

    doc.save(out_path)
    return out_path

def main():
    parser = argparse.ArgumentParser(description="Arabic poetry layout with optional merging/shift/equalize for two-row layout.")
    parser.add_argument("--in", dest="infile", default=None, help="UTF-8 file: each line is 'first|second' or Tab")
    parser.add_argument("--out", dest="outfile", default="poetry_aligned.docx", help="Output .docx filename")
    parser.add_argument("--page", choices=["A4","A5","Letter"], default="A4", help="Page size")
    parser.add_argument("--margin", type=float, default=2.0, help="Uniform page margins (cm)")
    parser.add_argument("--w1", type=float, default=7.5, help="Right column width (cm)")
    parser.add_argument("--w2", type=float, default=7.5, help="Left column width (cm)")
    parser.add_argument("--gap", type=float, default=0.0, help="Gap between hemistichs (cm)")
    parser.add_argument("--gap-mode", choices=["indent","column"], default="column", help="Gap as indent or third column")
    parser.add_argument("--font", type=str, default="Traditional Arabic", help="Font family")
    parser.add_argument("--size", type=float, default=18, help="Font size (pt)")
    parser.add_argument("--borders", action="store_true", help="Show table borders")
    parser.add_argument("--rtl-off", action="store_true", help="Disable RTL (default on)")
    parser.add_argument("--right-align", choices=["right","justify","distribute"], default="right", help="Alignment for first hemistich")
    parser.add_argument("--second-align", choices=["right","left","justify","distribute"], default="left", help="Alignment for second hemistich")
    parser.add_argument("--auto-widths", action="store_true", help="Compute w1/w2 from page width & margins automatically.")
    parser.add_argument("--layout", choices=["row","two-rows"], default="row", help="Place both hemistichs in the same row or separate rows.")
    parser.add_argument("--nowrap", action="store_true", help="Prevent line wrapping inside hemistich cells.")
    # الإضافات الجديدة (اختيارية بالكامل)
    parser.add_argument("--merge-second", action="store_true", help="ادمج خلايا الصف الثاني لكل بيت في خلية واحدة (two-rows فقط).")
    parser.add_argument("--second-shift-pct", type=float, default=0.0, help="إزاحة الشطر الثاني للداخل بنسبة من w1 (two-rows + merge-second).")
    parser.add_argument("--equalize-lengths", action="store_true", help="اجعل مساحة نص الشطر الثاني مساويةً لمساحة نص الشطر الأول (two-rows + merge-second).")

    args = parser.parse_args()

    # Load poems
    if args.infile:
        poems = read_poems_from_file(args.infile)
        if not poems:
            print("⚠️ الملف لا يحتوي أبياتًا صالحة.", file=sys.stderr); sys.exit(1)
    else:
        poems = [
            ("في جنة الخُلد ما لقانا وموعدنا", "وسوف أبقى لذاك اليوم أنتظرُ"),
            ("ها جئتَ يا إخوتي أهدي فؤادي لي", "أشعُّ بالنور من جبلٍ إلى جبل"),
        ]

    out = build_poetry_docx_align(
        poems,
        out_path=args.outfile,
        col_width_cm=(args.w1, args.w2),
        gap_cm=args.gap,
        gap_mode=args.gap_mode,
        font_name=args.font,
        font_size_pt=args.size,
        show_borders=args.borders,
        rtl=(not args.rtl_off),
        right_align_mode=args.right_align,
        second_align_mode=args.second_align,
        auto_widths=args.auto_widths,
        layout=args.layout,
        page=args.page,
        margin=args.margin,
        nowrap=args.nowrap,
        merge_second=args.merge_second,
        second_shift_pct=args.second_shift_pct,
        equalize_lengths=args.equalize_lengths,
    )
    print("✅ Saved ->", out)

if __name__ == "__main__":
    main()
