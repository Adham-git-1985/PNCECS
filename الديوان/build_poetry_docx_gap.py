# -*- coding: utf-8 -*-
"""
build_poetry_docx_align.py (auto-widths)
----------------------------------------
- يثبّت آخر حرف في الشطر الأول (right/distribute/justify).
- يدعم محاذاة الشطر الثاني (left/distribute/justify).
- فراغ بين الشطرين (--gap) بأسلوبين: --gap-mode column | indent.
- جديد: --auto-widths يحسب عرض العمودين تلقائيًا من عرض الصفحة وهوامشها:
    * في gap-mode=column: w1 = w2 = (available - gap) / 2
    * في gap-mode=indent : w1 = w2 = available / 2  (والـ gap يُستخدم كمسافات بادئة داخلية)

تشغيل مثال موصى به:
    python build_poetry_docx_align.py --out ديواني.docx --auto-widths --gap 1.0 --gap-mode column \
        --right-align distribute --second-align distribute
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys


# ----------------- Helpers -----------------
def set_paragraph_rtl(paragraph, enable=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for tag, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn('w:w'), str(int(val)))
        el.set(qn('w:type'), 'dxa')


def set_table_borders(table, show=False):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn('w:val'), 'single' if show else 'nil')


def set_table_direction_rtl(table, enable=True):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bidi = tblPr.find(qn('w:bidiVisual'))
    if bidi is None:
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')


def read_poems_from_file(path):
    poems = []
    for raw in Path(path).read_text(encoding='utf-8').splitlines():
        line = raw.strip()
        if not line:
            continue
        if '|' in line:
            a, b = line.split('|', 1)
        elif '\t' in line:
            a, b = line.split('\t', 1)
        else:
            a, b = line, ''
        poems.append((a.strip(), b.strip()))
    return poems


# ----------------- Alignment mapping -----------------
def map_align(value, default='right'):
    val = (value or default).lower()
    if val == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    if val == 'left':
        return WD_ALIGN_PARAGRAPH.LEFT
    if val == 'justify':
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if val == 'distribute':
        return WD_ALIGN_PARAGRAPH.DISTRIBUTE
    return WD_ALIGN_PARAGRAPH.RIGHT


# ----------------- Core -----------------
def compute_auto_widths(doc, gap_cm, gap_mode):
    """Return (w1_cm, w2_cm, available_cm) based on page width and margins."""
    sec = doc.sections[0]
    page_cm = sec.page_width.cm
    left_cm = sec.left_margin.cm
    right_cm = sec.right_margin.cm
    available = max(0.1, page_cm - left_cm - right_cm)
    if gap_mode == 'column':
        inner = max(0.1, available - max(0.0, gap_cm))
        each = max(0.1, inner / 2.0)
        return each, each, available
    else:
        # indent mode doesn't consume width; columns split the available space
        each = max(0.1, available / 2.0)
        return each, each, available


def build_poetry_docx_align(
    poems, out_path="poetry_aligned.docx",
    col_width_cm=(7.5, 7.5), gap_cm=0.0, gap_mode="column",
    font_name="Traditional Arabic", font_size_pt=18,
    show_borders=False, rtl=True,
    right_align_mode='right', second_align_mode='left',
    auto_widths=False
):
    """
    right_align_mode : محاذاة الشطر الأول  {right, justify, distribute}
    second_align_mode: محاذاة الشطر الثاني {left,  justify, distribute}
    auto_widths      : يحسب w1/w2 تلقائياً من عرض الصفحة وهوامشها.
    """
    if gap_mode not in ("indent", "column"):
        raise ValueError("gap_mode must be 'indent' or 'column'")

    doc = Document()

    # Auto widths if requested
    if auto_widths:
        w1, w2, avail = compute_auto_widths(doc, gap_cm, gap_mode)
    else:
        w1, w2 = col_width_cm



    align_first  = map_align(right_align_mode, default='right')
    align_second = map_align(second_align_mode, default='left')

    if gap_mode == "column":
        table = doc.add_table(rows=len(poems), cols=3)
        table.autofit = False
        table.columns[0].width = Cm(w1)
        table.columns[1].width = Cm(max(0.0, gap_cm))
        table.columns[2].width = Cm(w2)
        set_table_borders(table, show=show_borders)
        set_table_direction_rtl(table, True)

        for i, (right_hemistich, left_hemistich) in enumerate(poems):
            row = table.rows[i]

            # الشطر الأول (عمود 0 يمينًا بصريًا)
            c_r = row.cells[0]; c_r.width = Cm(w1)
            set_cell_margins(c_r, 0, 0, 0, 0)
            p_r = c_r.paragraphs[0]
            p_r.alignment = align_first
            set_paragraph_rtl(p_r, rtl)
            run_r = p_r.add_run(right_hemistich)
            run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

            # عمود الفاصل
            row.cells[1].width = Cm(max(0.0, gap_cm))

            # الشطر الثاني (عمود 2 يسارًا بصريًا)
            c_l = row.cells[2]; c_l.width = Cm(w2)
            set_cell_margins(c_l, 0, 0, 0, 0)
            p_l = c_l.paragraphs[0]
            p_l.alignment = align_second
            set_paragraph_rtl(p_l, rtl)
            run_l = p_l.add_run(left_hemistich)
            run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

    else:
        # indent mode
        table = doc.add_table(rows=len(poems), cols=2)
        table.autofit = False
        table.columns[0].width = Cm(w1)
        table.columns[1].width = Cm(w2)
        set_table_borders(table, show=show_borders)
        set_table_direction_rtl(table, True)

        for i, (right_hemistich, left_hemistich) in enumerate(poems):
            row = table.rows[i]

            c_r = row.cells[0]; c_r.width = Cm(w1)
            set_cell_margins(c_r, 0, 0, 0, 0)
            p_r = c_r.paragraphs[0]
            p_r.alignment = align_first
            set_paragraph_rtl(p_r, rtl)
            if gap_cm > 0:
                p_r.paragraph_format.right_indent = Cm(gap_cm)
            run_r = p_r.add_run(right_hemistich)
            run_r.font.name = font_name; run_r.font.size = Pt(font_size_pt)

            c_l = row.cells[1]; c_l.width = Cm(w2)
            set_cell_margins(c_l, 0, 0, 0, 0)
            p_l = c_l.paragraphs[0]
            p_l.alignment = align_second
            set_paragraph_rtl(p_l, rtl)
            if gap_cm > 0:
                p_l.paragraph_format.left_indent = Cm(gap_cm)
            run_l = p_l.add_run(left_hemistich)
            run_l.font.name = font_name; run_l.font.size = Pt(font_size_pt)

    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Align both hemistichs with optional gap and auto widths from page size.")
    parser.add_argument("--in", dest="infile", default=None, help="UTF-8 file: each line is a verse: first|second or Tab")
    parser.add_argument("--out", dest="outfile", default="poetry_aligned.docx", help="Output .docx filename")
    parser.add_argument("--w1", type=float, default=7.5, help="Right column width (cm)")
    parser.add_argument("--w2", type=float, default=7.5, help="Left column width (cm)")
    parser.add_argument("--gap", type=float, default=0.0, help="Gap between hemistichs (cm)")
    parser.add_argument("--gap-mode", choices=["indent","column"], default="column", help="Gap as indent or third column")
    parser.add_argument("--font", type=str, default="Traditional Arabic", help="Font family")
    parser.add_argument("--size", type=float, default=18, help="Font size (pt)")
    parser.add_argument("--borders", action="store_true", help="Show table borders")
    parser.add_argument("--rtl-off", action="store_true", help="Disable RTL (default on)")
    parser.add_argument("--right-align", choices=["right","justify","distribute"], default="right",
                        help="Alignment for first hemistich")
    parser.add_argument("--second-align", choices=["left","justify","distribute"], default="left",
                        help="Alignment for second hemistich")
    parser.add_argument("--auto-widths", action="store_true", help="Compute w1/w2 from page width & margins automatically.")
    args = parser.parse_args()

    # Load poems
    if args.infile:
        poems = read_poems_from_file(args.infile)
        if not poems:
            print("⚠️ الملف لا يحتوي أبياتًا صالحة.", file=sys.stderr); sys.exit(1)
    else:
        poems = [
            ("ها جئتَ يا إخوتي أهدي فؤادي لي", "أشعُّ بالنور من جبلٍ إلى جبل"),
            ("رسمتُ من كلماتِ الحبِّ صورتَها", "ولوَّحتُ لكم فرحى مناديل"),
            ("قد صغتُها من وفائي وعاطفتي", "وزيَّنتُ بفخّارٍ بالأكاليل"),
            ("يا صفوةَ الناس يا أحبابَ مسندي", "هَيّا ارفقوها وجودوا بالتهاليل"),
        ]

    out = build_poetry_docx_align(
        poems,
        out_path=args.outfile,
        col_width_cm=(args.w1, args.w2),
        gap_cm=args.gap,
        gap_mode=args.gap_mode,
        font_name=args.font,
        font_size_pt=args.size,
        show_borders=args.borders,
        rtl=(not args.rtl_off),
        right_align_mode=args.right_align,
        second_align_mode=args.second_align,
        auto_widths=args.auto_widths
    )
    print("✅ Saved ->", out)


if __name__ == "__main__":
    main()
