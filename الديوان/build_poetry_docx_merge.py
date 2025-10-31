# -*- coding: utf-8 -*-
"""
build_poetry_docx_merge.py
--------------------------
- تنسيق دواوين عربية في وورد.
- وضع two-rows: الشطر الأول في صف مستقل، والشطر الثاني في الصف التالي.
- خيار جديد: --merge-second لدمج خلايا الصف الثاني (الشطر الثاني) في خلية واحدة
  ثم محاذاتها يمينًا، مع ترك فراغ يسارها يساوي نسبة (--second-shift-pct) من عرض عمود الشطر الأول w1.
  (في العربية RTL يكون "END" = اليسار، لذا نستخدم هامش END للخلية المدموجة).

مثال التشغيل (A5 + gap بسيط + دمج صف الشطر الثاني + إزاحة 75%):
    python build_poetry_docx_merge.py --in poems.txt --out ديواني.docx --page A5 --margin 1 \
      --layout two-rows --auto-widths --gap 0.05 --gap-mode indent \
      --right-align distribute --second-align right \
      --font "Sakkal Majalla" --size 18 --nowrap \
      --merge-second --second-shift-pct 75
"""

import argparse
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys

# ----------------- Helpers -----------------
def set_paragraph_rtl(paragraph, enable=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for tag, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn('w:w'), str(int(val)))
        el.set(qn('w:type'), 'dxa')

def set_table_borders(table, show=False):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            borders.append(el)
        el.set(qn('w:val'), 'single' if show else 'nil')

def set_table_direction_rtl(table, enable=True):
    tblPr = table._tbl.tblPr
    bidi = tblPr.find(qn('w:bidiVisual'))
    if bidi is None:
        bidi = OxmlElement('w:bidiVisual')
        tblPr.append(bidi)
    bidi.set(qn('w:val'), '1' if enable else '0')

def disable_cell_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    no_wrap = tcPr.find(qn('w:noWrap'))
    if no_wrap is None:
        no_wrap = OxmlElement('w:noWrap')
        tcPr.append(no_wrap)

def read_poems_from_file(path):
    poems = []
    for raw in Path(path).read_text(encoding='utf-8').splitlines():
        line = raw.strip()
        if not line:
            continue
        if '|' in line:
            a, b = line.split('|', 1)
        elif '\t' in line:
            a, b = line.split('\t', 1)
        else:
            a, b = line, ''
        poems.append((a.strip(), b.strip()))
    return poems

def map_align(value, default='right'):
    val = (value or default).lower()
    if val == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    if val == 'left':
        return WD_ALIGN_PARAGRAPH.LEFT
    if val == 'justify':
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if val == 'distribute':
        return WD_ALIGN_PARAGRAPH.DISTRIBUTE
    return WD_ALIGN_PARAGRAPH.RIGHT

def apply_page_setup(doc, page, margin_cm):
    sec = doc.sections[0]
    page = (page or 'A4').upper()
    if page == 'A5':
        sec.page_width  = Cm(14.8)
        sec.page_height = Cm(21.0)
    elif page == 'LETTER':
        sec.page_width  = Cm(21.59); sec.page_height = Cm(27.94)
    else:  # A4
        sec.page_width  = Cm(21.0);  sec.page_height = Cm(29.7)
    if margin_cm is not None:
        m = Cm(margin_cm)
        sec.left_margin = m; sec.right_margin = m; sec.top_margin = m; sec.bottom_margin = m

def compute_auto_widths(doc, gap_cm, gap_mode):
    sec = doc.sections[0]
    available = sec.page_width.cm - sec.left_margin.cm - sec.right_margin.cm
    available = max(0.1, available)
    if gap_mode == 'column':
        inner = max(0.1, available - max(0.0, gap_cm))
        each = max(0.1, inner / 2.0)
        return each, each
    else:
        each = max(0.1, available / 2.0)
        return each, each

def cm_to_twips(cm):
    return int(round(cm * 566.9291339))  # 1cm ≈ 566.93 twips

# ----------------- Core -----------------
def build_poetry_docx_align(
    poems, out_path="poetry_aligned.docx",
    col_width_cm=(7.5, 7.5), gap_cm=0.0, gap_mode="column",
    font_name="Traditional Arabic", font_size_pt=18,
    show_borders=False, rtl=True,
    right_align_mode='right', second_align_mode='right',
    auto_widths=False, layout='row',
    page='A4', margin=2.0, nowrap=False,
    line_spacing=None, row_gap_cm=0.0,
    bold=False, merge_second=False, second_shift_pct=75.0
):
    if gap_mode not in ("indent", "column"):
        raise ValueError("gap_mode must be 'indent' or 'column'")
    if layout not in ("row", "two-rows"):
        raise ValueError("layout must be 'row' or 'two-rows'")

    doc = Document()
    apply_page_setup(doc, page, margin)

    # compute widths
    if auto_widths:
        w1, w2 = compute_auto_widths(doc, gap_cm, gap_mode)
    else:
        w1, w2 = col_width_cm

    align_first  = map_align(right_align_mode, default='right')
    align_second = map_align(second_align_mode, default='right')

    def apply_para_opts(p):
        if line_spacing:
            p.paragraph_format.line_spacing = float(line_spacing)

    shift_cm = max(0.0, (second_shift_pct or 0) * 0.01 * w1) if layout == 'two-rows' else 0.0

    if layout == "row":
        # لا حاجة للدمج هنا
        table = doc.add_table(rows=len(poems), cols=3 if gap_mode=="column" else 2)
        table.autofit = False
        if gap_mode == "column":
            table.columns[0].width = Cm(w1); table.columns[1].width = Cm(max(0.0, gap_cm)); table.columns[2].width = Cm(w2)
        else:
            table.columns[0].width = Cm(w1); table.columns[1].width = Cm(w2)
        set_table_borders(table, show=show_borders); set_table_direction_rtl(table, True)

        for i,(h1,h2) in enumerate(poems):
            r = table.rows[i]
            c1 = r.cells[0]; set_cell_margins(c1,0,0,0,0); 
            if nowrap: disable_cell_wrap(c1)
            p1 = c1.paragraphs[0]; p1.alignment=align_first; set_paragraph_rtl(p1, rtl); apply_para_opts(p1)
            run1 = p1.add_run(h1); run1.font.name=font_name; run1.font.size=Pt(font_size_pt); run1.font.bold=bold

            if gap_mode=="column":
                r.cells[1].width = Cm(max(0.0, gap_cm))
                c2 = r.cells[2]
            else:
                c2 = r.cells[1]
            set_cell_margins(c2,0,0,0,0)
            if nowrap: disable_cell_wrap(c2)
            p2 = c2.paragraphs[0]; p2.alignment=align_second; set_paragraph_rtl(p2, rtl); apply_para_opts(p2)
            run2 = p2.add_run(h2); run2.font.name=font_name; run2.font.size=Pt(font_size_pt); run2.font.bold=bold

    else:
        # two-rows
        cols = 3 if gap_mode=="column" else 2
        table = doc.add_table(rows=len(poems)*2, cols=cols)
        table.autofit = False
        if cols==3:
            table.columns[0].width = Cm(w1); table.columns[1].width = Cm(max(0.0, gap_cm)); table.columns[2].width = Cm(w2)
        else:
            table.columns[0].width = Cm(w1); table.columns[1].width = Cm(w2)
        set_table_borders(table, show=show_borders); set_table_direction_rtl(table, True)

        for i,(h1,h2) in enumerate(poems):
            r1 = table.rows[i*2]; r2 = table.rows[i*2+1]
            # الشطر الأول
            c1 = r1.cells[0]; set_cell_margins(c1,0,0,0,0)
            if nowrap: disable_cell_wrap(c1)
            p1 = c1.paragraphs[0]; p1.alignment=align_first; set_paragraph_rtl(p1, rtl); apply_para_opts(p1)
            run1 = p1.add_run(h1); run1.font.name=font_name; run1.font.size=Pt(font_size_pt); run1.font.bold=bold

            # الشطر الثاني
            if not merge_second:
                # بدون دمج: يبقى كما هو (للمرجعية)
                if cols==3:
                    r1.cells[1].width = Cm(max(0.0, gap_cm))
                    c2 = r2.cells[2]; set_cell_margins(c2,0,0,0,0)
                else:
                    c2 = r2.cells[1]; set_cell_margins(c2,0,0,0,0)
                if nowrap: disable_cell_wrap(c2)
                p2 = c2.paragraphs[0]; p2.alignment=align_second; set_paragraph_rtl(p2, rtl); apply_para_opts(p2)
                run2 = p2.add_run(h2); run2.font.name=font_name; run2.font.size=Pt(font_size_pt); run2.font.bold=bold
            else:
                # دمج خلايا الصف الثاني كلها ثم محاذاة يمين + هامش END = 75% من w1 (أو حسب النسبة)
                if cols==3:
                    merged = r2.cells[0].merge(r2.cells[1]).merge(r2.cells[2])
                else:
                    merged = r2.cells[0].merge(r2.cells[1])
                # هامش END (يسار في RTL) = shift_cm
                set_cell_margins(merged, 0, 0, 0, cm_to_twips(shift_cm))
                if nowrap: disable_cell_wrap(merged)
                p2 = merged.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_paragraph_rtl(p2, rtl); apply_para_opts(p2)
                run2 = p2.add_run(h2); run2.font.name=font_name; run2.font.size=Pt(font_size_pt); run2.font.bold=bold

    doc.save(out_path)
    return out_path

def main():
    parser = argparse.ArgumentParser(description="Arabic poetry layout with optional merging of second row and inward shift by % of w1.")
    parser.add_argument("--in", dest="infile", default=None, help="UTF-8 file: each line = first|second (or Tab).")
    parser.add_argument("--out", dest="outfile", default="poetry_aligned.docx", help="Output .docx filename.")
    parser.add_argument("--page", choices=["A4","A5","Letter"], default="A4", help="Page size.")
    parser.add_argument("--margin", type=float, default=2.0, help="Uniform margins (cm).")
    parser.add_argument("--w1", type=float, default=7.5, help="Right column width (cm).")
    parser.add_argument("--w2", type=float, default=7.5, help="Left column width (cm).")
    parser.add_argument("--gap", type=float, default=0.0, help="Gap between hemistichs (cm).")
    parser.add_argument("--gap-mode", choices=["indent","column"], default="column", help="Gap as indent or third column.")
    parser.add_argument("--font", type=str, default="Traditional Arabic", help="Font family.")
    parser.add_argument("--size", type=float, default=18, help="Font size (pt).")
    parser.add_argument("--borders", action="store_true", help="Show table borders.")
    parser.add_argument("--rtl-off", action="store_true", help="Disable RTL (default on).")
    parser.add_argument("--right-align", choices=["right","justify","distribute"], default="right", help="Alignment for first hemistich.")
    parser.add_argument("--second-align", choices=["right","left","justify","distribute"], default="right", help="Alignment for second hemistich.")
    parser.add_argument("--auto-widths", action="store_true", help="Compute w1/w2 from page width & margins automatically.")
    parser.add_argument("--layout", choices=["row","two-rows"], default="row", help="Same row or each in its own row.")
    parser.add_argument("--nowrap", action="store_true", help="Prevent line wrapping inside hemistich cells.")
    parser.add_argument("--line-spacing", type=float, default=None, help="Paragraph line spacing multiplier (e.g., 1.0, 1.15, 1.5, 2.0).")
    parser.add_argument("--row-gap", type=float, default=0.0, help="Vertical space (cm) between shatr1 and shatr2 in two-rows layout.")
    parser.add_argument("--bold", action="store_true", help="Make poem text bold.")
    parser.add_argument("--merge-second", action="store_true", help="Merge all cells of the second row of each verse into one cell.")
    parser.add_argument("--second-shift-pct", type=float, default=75.0, help="White space on the LEFT of merged second row as % of w1 (two-rows only).")
    args = parser.parse_args()

    # Load poems
    if args.infile:
        poems = read_poems_from_file(args.infile)
        if not poems:
            print("⚠️ الملف لا يحتوي أبياتًا صالحة.", file=sys.stderr); sys.exit(1)
    else:
        poems = [
            ("في جنة الخُلد ما لقانا وموعدنا", "وسوف أبقى لذاك اليوم أنتظرُ"),
            ("ها جئتَ يا إخوتي أهدي فؤادي لي", "أشعُّ بالنور من جبلٍ إلى جبل"),
        ]

    out = build_poetry_docx_align(
        poems,
        out_path=args.outfile,
        col_width_cm=(args.w1, args.w2),
        gap_cm=args.gap,
        gap_mode=args.gap_mode,
        font_name=args.font,
        font_size_pt=args.size,
        show_borders=args.borders,
        rtl=(not args.rtl_off),
        right_align_mode=args.right_align,
        second_align_mode=args.second_align,
        auto_widths=args.auto_widths,
        layout=args.layout,
        page=args.page,
        margin=args.margin,
        nowrap=args.nowrap,
        line_spacing=args.line_spacing,
        row_gap_cm=args.row_gap,
        bold=args.bold,
        merge_second=args.merge_second,
        second_shift_pct=args.second_shift_pct
    )
    print("✅ Saved ->", out)

if __name__ == "__main__":
    main()
